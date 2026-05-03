"""Convert jpt_draft.md to a Word document with charts embedded inline.

Outputs to the user's Windows Downloads folder so they can open it directly
from Windows, edit, and upload back.
"""
import re
from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

HERE     = Path(__file__).resolve().parent
CHARTS   = HERE / 'charts'
SRC      = HERE / 'jpt_draft.md'
OUT      = Path('/mnt/c/Users/Pushpesh/Downloads/jpt_draft.docx')

# Chart placeholder text → image filename
CHART_IMAGES = {
    'CHART 1': CHARTS / 'chart1_perba_p90_floor.png',
    'CHART 2': CHARTS / 'chart2_pjm_capacity_auction.png',
    'CHART 3': CHARTS / 'chart3_gevernova_slots.png',
    'CHART 4': CHARTS / 'chart4_load_queue_waterfall.png',
}

INLINE_RE = re.compile(r'(\*\*[^*]+\*\*|\*[^*]+\*|\[\^\d+\])')


def add_runs(paragraph, text):
    """Render inline markdown (**bold**, *italic*, [^N] footnotes) into runs."""
    for token in INLINE_RE.split(text):
        if not token:
            continue
        if token.startswith('**') and token.endswith('**'):
            r = paragraph.add_run(token[2:-2])
            r.bold = True
        elif token.startswith('*') and token.endswith('*') and len(token) > 2:
            r = paragraph.add_run(token[1:-1])
            r.italic = True
        elif token.startswith('[^') and token.endswith(']'):
            num = token[2:-1]
            r = paragraph.add_run(num)
            r.font.superscript = True
            r.font.size = Pt(8)
            r.font.color.rgb = RGBColor(0xC0, 0x39, 0x2B)
        else:
            paragraph.add_run(token)


def add_image(doc, path, caption=None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(path), width=Inches(6.4))
    if caption:
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = cap.add_run(caption)
        r.italic = True
        r.font.size = Pt(9)
        r.font.color.rgb = RGBColor(0x7F, 0x8C, 0x8D)


def main():
    text = SRC.read_text()

    # Strip the existing markdown footnote DEFINITIONS — we will render them
    # below as a separate Sources section. Keep the [^N] reference markers in
    # the body.
    body, _, footnote_block = text.partition('## Footnotes')

    # Parse footnote definitions: [^N]: ...
    fn_def_re = re.compile(r'^\[\^(\d+)\]:\s*(.+?)(?=^\[\^|\Z)', re.M | re.S)
    footnotes = {m.group(1): re.sub(r'\s+', ' ', m.group(2).strip())
                 for m in fn_def_re.finditer(footnote_block)}

    doc = Document()

    # Margins
    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    # Default body font
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    lines = body.splitlines()
    i = 0
    in_chart_placeholder = False
    pending_chart_key = None

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # End of chart placeholder block (caption italic line right after)
        if in_chart_placeholder:
            if stripped.startswith('> *') and pending_chart_key:
                caption = stripped.lstrip('> ').strip('*')
                img_path = CHART_IMAGES.get(pending_chart_key)
                if img_path and img_path.exists():
                    add_image(doc, img_path, caption=caption)
                in_chart_placeholder = False
                pending_chart_key = None
                i += 1
                continue
            elif stripped.startswith('>'):
                i += 1
                continue
            else:
                # Block ended without caption
                if pending_chart_key:
                    img_path = CHART_IMAGES.get(pending_chart_key)
                    if img_path and img_path.exists():
                        add_image(doc, img_path)
                in_chart_placeholder = False
                pending_chart_key = None
                # Fall through to normal processing of current line

        # Detect chart placeholder start: > **[CHART N — ...]**
        m = re.match(r'>\s*\*\*\[(CHART \d+)\b', stripped)
        if m:
            in_chart_placeholder = True
            pending_chart_key = m.group(1)
            i += 1
            continue

        # H1 — title line
        if stripped.startswith('# ') and not stripped.startswith('##'):
            heading = stripped[2:].strip()
            p = doc.add_paragraph()
            r = p.add_run(heading)
            r.bold = True
            r.font.size = Pt(20)
            r.font.color.rgb = RGBColor(0x2C, 0x3E, 0x50)
            i += 1
            continue

        # H2 — section heading (and HOOK markers)
        if stripped.startswith('## '):
            heading = stripped[3:].strip()
            p = doc.add_paragraph()
            r = p.add_run(heading)
            r.bold = True
            r.font.size = Pt(15)
            r.font.color.rgb = RGBColor(0xC0, 0x39, 0x2B) if heading.startswith('[HOOK') else RGBColor(0x2C, 0x3E, 0x50)
            i += 1
            continue

        # Horizontal rule
        if stripped == '---':
            p = doc.add_paragraph()
            r = p.add_run('—' * 30)
            r.font.color.rgb = RGBColor(0xBD, 0xC3, 0xC7)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            i += 1
            continue

        # Numbered list item
        m = re.match(r'^(\d+)\.\s+(.+)', stripped)
        if m:
            p = doc.add_paragraph(style='List Number')
            add_runs(p, m.group(2))
            i += 1
            continue

        # Bullet
        if stripped.startswith('- '):
            p = doc.add_paragraph(style='List Bullet')
            add_runs(p, stripped[2:])
            i += 1
            continue

        # Italic-only "draft metadata" line at top (single-line *...*)
        if stripped.startswith('*') and stripped.endswith('*') and not stripped.startswith('**') \
                and stripped.count('*') == 2:
            p = doc.add_paragraph()
            r = p.add_run(stripped.strip('*'))
            r.italic = True
            r.font.color.rgb = RGBColor(0x7F, 0x8C, 0x8D)
            i += 1
            continue

        # Regular paragraph — collect until blank line
        if stripped:
            buf = [stripped]
            j = i + 1
            while j < len(lines) and lines[j].strip() and not re.match(
                    r'^(#{1,6}\s|\d+\.\s|-\s|>|---)', lines[j].strip()):
                buf.append(lines[j].strip())
                j += 1
            p = doc.add_paragraph()
            add_runs(p, ' '.join(buf))
            i = j
            continue

        i += 1

    # Sources section
    doc.add_page_break()
    p = doc.add_paragraph()
    r = p.add_run('Sources')
    r.bold = True
    r.font.size = Pt(15)
    r.font.color.rgb = RGBColor(0x2C, 0x3E, 0x50)

    for num in sorted(footnotes, key=int):
        p = doc.add_paragraph()
        r_num = p.add_run(f'{num}. ')
        r_num.bold = True
        r_num.font.color.rgb = RGBColor(0xC0, 0x39, 0x2B)
        # Linkify URLs, keep prose plain
        text_remaining = footnotes[num]
        for token in re.split(r'(https?://\S+)', text_remaining):
            if token.startswith('http'):
                r = p.add_run(token)
                r.font.color.rgb = RGBColor(0x1A, 0x73, 0xE8)
                r.font.underline = True
            else:
                p.add_run(token)
        p.paragraph_format.space_after = Pt(4)

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(f'wrote {OUT} ({OUT.stat().st_size:,} bytes)')


if __name__ == '__main__':
    main()
